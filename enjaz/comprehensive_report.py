"""
Comprehensive Analysis Report Module
Creates detailed analytical reports with student-subject breakdown
"""

import pandas as pd
from typing import List, Dict
from enjaz.analysis import get_band, get_band_emoji
from enjaz.parent_recommendations import get_parent_recommendation


def create_comprehensive_report(all_data: List[Dict]) -> pd.DataFrame:
    """
    Create comprehensive report with each row showing:
    - Student name
    - Grade
    - Section
    - Subject
    - Total assessments for subject
    - Completed assessments for subject
    - Completion rate for subject
    - Overall completion rate (across all subjects)
    - Band category
    - Recommendation
    
    Args:
        all_data: List of sheet data dictionaries
    
    Returns:
        DataFrame with comprehensive report
    """
    
    # Collect all student-subject combinations
    report_rows = []
    
    # First, collect all data for each student
    student_data = {}
    
    for sheet_data in all_data:
        subject = sheet_data.get('subject', sheet_data.get('sheet_name', 'غير محدد'))
        grade = sheet_data.get('grade', '')
        section = sheet_data.get('section', '')
        
        for student in sheet_data['students']:
            student_name = student['student_name']
            
            if student_name not in student_data:
                student_data[student_name] = {
                    'grade': grade,
                    'section': section,
                    'subjects': []
                }
            
            # Add subject data
            if student['has_due']:
                student_data[student_name]['subjects'].append({
                    'subject': subject,
                    'total_due': student['total_due'],
                    'completed': student['completed'],
                    'completion_rate': student['completion_rate']
                })
    
    # Now create rows for each student-subject combination
    for student_name, data in student_data.items():
        if not data['subjects']:
            continue
        
        # Calculate overall completion rate
        total_due_all = sum(s['total_due'] for s in data['subjects'])
        total_completed_all = sum(s['completed'] for s in data['subjects'])
        overall_rate = 100 * total_completed_all / total_due_all if total_due_all > 0 else 0
        
        # Get band and recommendation
        band = get_band(overall_rate)
        emoji = get_band_emoji(overall_rate)
        recommendation = get_parent_recommendation(overall_rate)
        
        # Create a row for each subject
        for subject_info in data['subjects']:
            report_rows.append({
                'اسم الطالب': student_name,
                'الصف': data['grade'],
                'الشعبة': data['section'],
                'المادة': subject_info['subject'],
                'إجمالي المادة': subject_info['total_due'],
                'منجز في المادة': subject_info['completed'],
                'نسبة الإنجاز للمادة (%)': round(subject_info['completion_rate'], 1),
                'النسبة الكلية للإنجاز (%)': round(overall_rate, 1),
                'الفئة': f"{emoji} {band}",
                'التوصية': recommendation
            })
    
    # Create DataFrame
    df = pd.DataFrame(report_rows)
    
    # Sort by student name, then by subject
    if not df.empty:
        df = df.sort_values(['اسم الطالب', 'المادة'])
        df = df.reset_index(drop=True)
    
    return df


def export_comprehensive_report_to_excel(df: pd.DataFrame, output_path: str, school_info: Dict = None):
    """
    Export comprehensive report to Excel with professional formatting.
    
    Args:
        df: Comprehensive report DataFrame
        output_path: Path to save Excel file
        school_info: Dictionary containing school information
    
    Returns:
        str: Path to saved file
    """
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.drawing.image import Image as XLImage
    from pathlib import Path
    from datetime import datetime
    
    if school_info is None:
        from enjaz.school_info import load_school_info
        school_info = load_school_info()
    
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Write data starting from row 10 to leave space for header
        df.to_excel(writer, sheet_name='التقرير التحليلي الشامل', index=False, startrow=9)
        
        # Get the worksheet
        worksheet = writer.sheets['التقرير التحليلي الشامل']
        
        # Add school header information
        header_font = Font(name='Arial', size=14, bold=True)
        normal_font = Font(name='Arial', size=11)
        center_alignment = Alignment(horizontal='center', vertical='center')
        right_alignment = Alignment(horizontal='right', vertical='center')
        
        # Add ministry logo if exists
        assets_path = Path(__file__).parent / 'assets'
        logo_path = assets_path / 'ministry_logo.png'
        if logo_path.exists():
            try:
                img = XLImage(str(logo_path))
                img.width = 80
                img.height = 80
                worksheet.add_image(img, 'A1')
            except:
                pass
        
        # School name
        worksheet['E1'] = school_info.get('school_name', '')
        worksheet['E1'].font = Font(name='Arial', size=16, bold=True)
        worksheet['E1'].alignment = center_alignment
        
        # Report title
        worksheet['E2'] = 'التقرير التحليلي الشامل للتقييمات الأسبوعية'
        worksheet['E2'].font = Font(name='Arial', size=14, bold=True)
        worksheet['E2'].alignment = center_alignment
        
        # Date
        worksheet['E3'] = f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}"
        worksheet['E3'].font = normal_font
        worksheet['E3'].alignment = center_alignment
        
        # School leadership information
        row = 5
        leadership = [
            ('مدير المدرسة', school_info.get('principal', '')),
            ('النائب الأكاديمي', school_info.get('academic_deputy', '')),
            ('النائب الإداري', school_info.get('admin_deputy', '')),
            ('منسق المشاريع', school_info.get('projects_coordinator', ''))
        ]
        
        for title, name in leadership:
            if name:
                worksheet[f'B{row}'] = f"{title}:"
                worksheet[f'B{row}'].font = Font(name='Arial', size=10, bold=True)
                worksheet[f'B{row}'].alignment = right_alignment
                
                worksheet[f'C{row}'] = name
                worksheet[f'C{row}'].font = Font(name='Arial', size=10)
                worksheet[f'C{row}'].alignment = right_alignment
                
                row += 1
        
        # Format header row (row 10)
        header_fill = PatternFill(start_color='8A1538', end_color='8A1538', fill_type='solid')
        header_font_white = Font(name='Arial', size=11, bold=True, color='FFFFFF')
        
        for cell in worksheet[10]:
            cell.fill = header_fill
            cell.font = header_font_white
            cell.alignment = center_alignment
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        
        # Format data cells
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in worksheet.iter_rows(min_row=11, max_row=worksheet.max_row, min_col=1, max_col=worksheet.max_column):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='right', vertical='center')
                cell.font = Font(name='Arial', size=10)
        
        # Auto-adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Freeze panes (header row)
        worksheet.freeze_panes = 'A11'
    
    return output_path


def export_comprehensive_report_to_word(df: pd.DataFrame, output_path: str, school_info: Dict = None):
    """
    Export comprehensive report to Word document.
    
    Args:
        df: Comprehensive report DataFrame
        output_path: Path to save Word file
        school_info: Dictionary containing school information
    
    Returns:
        str: Path to saved file
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    from pathlib import Path
    
    if school_info is None:
        from enjaz.school_info import load_school_info
        school_info = load_school_info()
    
    doc = Document()
    
    # Set RTL for the document
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    
    # Add logo if exists
    assets_path = Path(__file__).parent / 'assets'
    logo_path = assets_path / 'ministry_logo.png'
    if logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(1.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    # School name
    school_name = doc.add_paragraph(school_info.get('school_name', ''))
    school_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_name.runs[0].font.size = Pt(18)
    school_name.runs[0].font.bold = True
    
    # Report title
    title = doc.add_paragraph('التقرير التحليلي الشامل للتقييمات الأسبوعية')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # Date
    date_para = doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Leadership info
    doc.add_paragraph()
    leadership = [
        ('مدير المدرسة', school_info.get('principal', '')),
        ('النائب الأكاديمي', school_info.get('academic_deputy', '')),
        ('النائب الإداري', school_info.get('admin_deputy', '')),
        ('منسق المشاريع', school_info.get('projects_coordinator', ''))
    ]
    
    for title_text, name in leadership:
        if name:
            para = doc.add_paragraph(f"{title_text}: {name}")
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Add table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        header_cells[i].text = column_name
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(output_path)
    return output_path

